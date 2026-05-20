"""
毕业论文开题报告格式自动校验工具 (第一阶段 - 纯 Python 基础版)

功能：读取 .docx 文件，根据预设格式规则进行校验，生成精美的 HTML 诊断报告并自动在浏览器中打开。
依赖：python-docx (外部), re/sys/json/webbrowser (标准库)
用法：python Format_verify_tool.py <path_to_docx>
"""

import re
import sys
import webbrowser
import os
from datetime import datetime
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn


# ─────────────────────────────────────────────────────────────
# 数据结构
# ─────────────────────────────────────────────────────────────

@dataclass
class CheckResult:
    """单条校验结果"""
    name: str
    passed: bool
    message: str


@dataclass
class AnalysisReport:
    """最终校验报告"""
    passed_items: List[Dict[str, str]] = field(default_factory=list)
    failed_items: List[Dict[str, str]] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)

    def add_pass(self, name: str, message: str):
        self.passed_items.append({"name": name, "message": message})

    def add_fail(self, name: str, message: str, context_text: str = None):
        item = {"name": name, "message": message}
        if context_text:
            item["context_text"] = context_text
        self.failed_items.append(item)

    def add_error(self, message: str):
        self.errors.append(message)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "passed_items": self.passed_items,
            "failed_items": self.failed_items,
            "errors": self.errors,
            "summary": {
                "total_checks": len(self.passed_items) + len(self.failed_items),
                "passed": len(self.passed_items),
                "failed": len(self.failed_items),
                "errors": len(self.errors),
            },
        }


# ─────────────────────────────────────────────────────────────
# DocumentParser —— 负责读取文档、按区域切块
# ─────────────────────────────────────────────────────────────

class DocumentParser:
    """
    职责：读取 .docx 文件，将文档按逻辑区域切块，
    提供段落级元数据（字体、字号、对齐、行距等）供 Validator 使用。
    """

    # 一级标题正则（如 "1. 文献综述"、"2. 研究的主要内容"）
    RE_H1 = re.compile(r"^\s*\d+\.\s*.+")
    # 二级标题正则（如 "1.1 选题目的和意义"）
    RE_H2 = re.compile(r"^\s*\d+\.\d+\s*.+")
    # 参考文献条目正则
    RE_REF = re.compile(r"\[(\d+)\]")
    # 日期正则：YYYY-MM-DD 或 YYYY年MM月DD日
    RE_DATE = re.compile(
        r"(\d{4})\s*[-年]\s*(\d{1,2})\s*[-月]\s*(\d{1,2})\s*日?"
    )

    # 中国字号 → 磅值映射
    FONT_SIZE_MAP = {
        "一号": 36.0,
        "二号": 22.0,
        "三号": 16.0,
        "四号": 14.0,
        "小四": 12.0,
        "五号": 10.5,
        "小五": 9.0,
    }

    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc: Optional[Document] = None
        self.paragraphs: List[Any] = []

    def load(self) -> "DocumentParser":
        """加载文档，缓存段落列表"""
        self.doc = Document(self.file_path)
        self.paragraphs = list(self.doc.paragraphs)
        return self

    # ── 段落元数据提取 ──

    @staticmethod
    def get_paragraph_font_name(para) -> Optional[str]:
        """获取段落级别设定的字体名（取第一个 run）"""
        for run in para.runs:
            if run.font.name:
                return run.font.name
        return None

    @staticmethod
    def get_run_east_asian_font(run) -> Optional[str]:
        """
        获取 run 的东亚字体名（优先级高于 font.name）。
        用于检测宋体、黑体等中文字体设置。
        """
        try:
            rPr = run._element.rPr
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    ea = rFonts.get(qn('w:eastAsia'))
                    if ea:
                        return ea
        except Exception:
            pass
        return None

    @staticmethod
    def get_paragraph_effective_font(para) -> Optional[str]:
        """
        获取段落的有效字体名（优先取东亚字体，其次取 font.name）。
        """
        for run in para.runs:
            ea_font = DocumentParser.get_run_east_asian_font(run)
            if ea_font:
                return ea_font
            if run.font.name:
                return run.font.name
        return None

    @staticmethod
    def get_paragraph_font_size(para) -> Optional[float]:
        """获取段落字号（单位：磅 pt）"""
        for run in para.runs:
            if run.font.size:
                return run.font.size.pt
        return None

    @staticmethod
    def get_paragraph_bold(para) -> Optional[bool]:
        """获取段落是否加粗"""
        for run in para.runs:
            if run.bold is not None:
                return run.bold
        return None

    @staticmethod
    def get_paragraph_alignment(para) -> Optional[str]:
        """获取段落对齐方式，返回字符串标识"""
        align = para.alignment
        if align is None:
            return "INHERITED"
        mapping = {
            0: "LEFT",
            1: "CENTER",
            2: "RIGHT",
            3: "JUSTIFY",
        }
        return mapping.get(align, str(align))

    @staticmethod
    def get_paragraph_line_spacing(para) -> Optional[float]:
        """
        获取段落行距。
        返回 1.5 表示 1.5 倍行距，None 表示未设定。
        """
        pf = para.paragraph_format
        if pf.line_spacing_rule == WD_LINE_SPACING.MULTIPLE and pf.line_spacing:
            return float(pf.line_spacing)
        if pf.line_spacing_rule == WD_LINE_SPACING.EXACTLY and pf.line_spacing:
            return pf.line_spacing.pt / 12.0
        return None

    # ── 区域切块 ──

    def get_cover_paragraphs(self) -> List[Any]:
        """提取封面段落：第一个一级标题之前的全部段落"""
        for i, p in enumerate(self.paragraphs):
            if self.RE_H1.match(p.text.strip()):
                return self.paragraphs[:i]
        return self.paragraphs

    def find_section_text(self, start_pattern: str, end_pattern: str = "") -> str:
        """
        提取从匹配 start_pattern 的段落到匹配 end_pattern 的段落之间的全部文本。
        end_pattern 为空时，截取到下一个一级标题。
        """
        start_idx = None
        end_idx = len(self.paragraphs)
        for i, p in enumerate(self.paragraphs):
            text = p.text.strip()
            if start_idx is None and start_pattern in text:
                start_idx = i
                continue
            if start_idx is not None:
                if end_pattern and end_pattern in text:
                    end_idx = i
                    break
                if not end_pattern and self.RE_H1.match(text) and i > start_idx:
                    end_idx = i
                    break
        if start_idx is None:
            return ""
        return "\n".join(p.text for p in self.paragraphs[start_idx:end_idx])

    def get_reference_paragraphs(self) -> List[str]:
        """
        截取「参考文献」标题到下一个一级标题之间的所有段落文本，
        按 [n] 序号分割成独立文献条目列表。
        """
        ref_text = self.find_section_text("参考文献")
        if not ref_text:
            return []
        lines = ref_text.split("\n")
        if lines and "参考文献" in lines[0]:
            lines = lines[1:]
        raw = "\n".join(lines)
        parts = self.RE_REF.split(raw)
        refs = []
        for i in range(1, len(parts), 2):
            if i + 1 < len(parts):
                refs.append(parts[i + 1].strip())
        return refs

    def count_section_words(self, start_pattern: str, end_pattern: str = "") -> int:
        """统计指定区域的中文字数（去除空格、标点、代码等干扰字符后）"""
        text = self.find_section_text(start_pattern, end_pattern)
        if not text:
            return 0
        lines = text.split("\n")
        if lines and start_pattern in lines[0]:
            text = "\n".join(lines[1:])
        chinese_chars = re.findall(r"[一-鿿]", text)
        english_words = re.findall(r"[a-zA-Z]+", text)
        return len(chinese_chars) + len(english_words)

    def get_all_body_paragraphs(self) -> List[Any]:
        """获取正文段落（排除封面）"""
        cover = self.get_cover_paragraphs()
        return self.paragraphs[len(cover):]

    def find_h2_titles(self) -> List[str]:
        """提取所有二级标题文本"""
        return [p.text.strip() for p in self.paragraphs if self.RE_H2.match(p.text.strip())]

    def extract_body_citation_indices(self) -> set:
        """
        遍历正文所有段落（排除「参考文献」及之后的内容），
        提取正文中所有 [n] 引用标记的数字，返回 int 集合。
        """
        citation_re = re.compile(r"\[(\d+)\]")
        indices = set()
        in_refs = False
        for p in self.paragraphs:
            text = p.text.strip()
            # 一旦遇到「参考文献」标题，停止扫描正文
            if "参考文献" in text and not self.RE_H2.match(text):
                in_refs = True
                continue
            if in_refs:
                continue
            for m in citation_re.finditer(text):
                indices.add(int(m.group(1)))
        return indices

    def extract_ref_list_indices(self) -> set:
        """
        从参考文献列表中提取所有文献序号，返回 int 集合。
        例如 [1], [2], [3] → {1, 2, 3}
        """
        indices = set()
        ref_text = self.find_section_text("参考文献")
        if not ref_text:
            return indices
        for m in self.RE_REF.finditer(ref_text):
            indices.add(int(m.group(1)))
        return indices

    def find_paragraph_containing(self, keyword: str) -> Optional[Any]:
        """查找包含指定关键字的段落"""
        for p in self.paragraphs:
            if keyword in p.text:
                return p
        return None

    def find_paragraph_index(self, keyword: str) -> Optional[int]:
        """查找包含指定关键字的段落索引"""
        for i, p in enumerate(self.paragraphs):
            if keyword in p.text:
                return i
        return None


# ─────────────────────────────────────────────────────────────
# Validator —— 执行各项格式校验逻辑
# ─────────────────────────────────────────────────────────────

class Validator:
    """
    职责：基于 DocumentParser 提供的数据，执行全部校验规则，
    将结果写入 AnalysisReport。
    """

    RECENT_YEAR_THRESHOLD = 2024
    FOREIGN_LANG_RATIO = 0.50
    THESIS_TAG = "[D]"

    def __init__(self, parser: DocumentParser, config: dict = None):
        self.parser = parser
        self.report = AnalysisReport()
        self.config = config or {}

    def run_all(self) -> Dict[str, Any]:
        """执行全部校验，返回结果字典"""
        self._check_global_line_spacing()
        self._check_cover_main_title()
        self._check_cover_info_fields()

        # 条件执行：导师全角空格检查（默认开启）
        if self.config.get("check_tutor_space", True):
            self._check_cover_advisor_format()

        self._check_cover_date()
        self._check_h1_structure()
        self._check_h2_structure()
        self._check_literature_review_word_count()
        self._check_references()
        self._check_supervisor_opinion()
        self._check_citation_consistency()
        self._check_ref_preceding_blank()

        # 条件执行：参考文献悬挂缩进检查（默认开启）
        if self.config.get("check_indent", True):
            self._check_ref_hanging_indent()

        # 条件执行：进度安排时间线顺叙检查（默认开启）
        if self.config.get("check_timeline", True):
            self._check_schedule_timeline()

        self._check_dual_advisor_order()
        self._check_fonts_and_sizes()
        return self.report.to_dict()

    # ── 1. 全局行距校验 ──

    def _check_global_line_spacing(self):
        body_paras = self.parser.get_all_body_paragraphs()
        if not body_paras:
            self.report.add_error("无法获取正文段落，跳过行距校验")
            return

        target = self.config.get("line_spacing", 1.5)

        non_compliant = []
        bad_snippets = []
        for i, p in enumerate(body_paras):
            spacing = self.parser.get_paragraph_line_spacing(p)
            # 未设定行距（None）或显式设为非目标值，均判定为不合规
            if spacing is None or abs(spacing - target) > 0.05:
                non_compliant.append(i)
                snippet = p.text[:50] + ("..." if len(p.text) > 50 else "")
                if snippet.strip():
                    actual = f"{spacing}" if spacing is not None else "未设定"
                    bad_snippets.append(f"[段落{i}] {snippet} (行距: {actual})")

        if not non_compliant:
            self.report.add_pass("全局行距", f"正文行距统一为 {target} 倍行距")
        else:
            self.report.add_fail(
                "全局行距",
                f"共 {len(non_compliant)} 个段落行距不合规（未设定或非 {target} 倍）",
                context_text="\n".join(bad_snippets[:10]) if bad_snippets else None
            )

    # ── 2. 封面 —— 主标题校验 ──

    def _check_cover_main_title(self):
        cover_paras = self.parser.get_cover_paragraphs()
        if not cover_paras:
            self.report.add_error("无法提取封面内容")
            return

        target_text = "毕业论文(设计)开题报告"
        target_para = None
        for p in cover_paras:
            clean = p.text.replace("（", "(").replace("）", ")").strip()
            if target_text in clean or "毕业论文" in clean and "开题报告" in clean:
                target_para = p
                break

        if target_para is None:
            self.report.add_fail("封面主标题", "未找到包含「毕业论文(设计)开题报告」的段落")
            return

        font_name = self.parser.get_paragraph_font_name(target_para)
        font_size = self.parser.get_paragraph_font_size(target_para)
        bold = self.parser.get_paragraph_bold(target_para)

        issues = []
        if font_name and font_name not in ("SimSun", "宋体", "simsun"):
            issues.append(f"字体应为宋体，实际为 {font_name}")
        if font_size and abs(font_size - 36.0) > 1.0:
            issues.append(f"字号应为一号(36pt)，实际为 {font_size}pt")
        if bold is not None and not bold:
            issues.append("应为加粗，实际为非加粗")

        if not issues:
            self.report.add_pass("封面主标题", "主标题格式符合要求（宋体、一号、加粗）")
        else:
            self.report.add_fail("封面主标题", "；".join(issues))

    # ── 3. 封面 —— 信息栏校验 ──

    def _check_cover_info_fields(self):
        cover_paras = self.parser.get_cover_paragraphs()
        info_keywords = ["题目", "学院", "专业", "班级", "学号", "姓名", "导师"]
        info_paras = []
        for p in cover_paras:
            text = p.text.strip()
            for kw in info_keywords:
                if kw in text and len(text) < 80:
                    info_paras.append((kw, p))
                    break

        if not info_paras:
            self.report.add_error("封面中未找到信息栏（题目/学院/专业等）")
            return

        issues = []
        for kw, p in info_paras:
            font_size = self.parser.get_paragraph_font_size(p)
            font_name = self.parser.get_paragraph_font_name(p)
            bold = self.parser.get_paragraph_bold(p)
            alignment = self.parser.get_paragraph_alignment(p)

            if font_size and abs(font_size - 14.0) > 1.0:
                issues.append(f"「{kw}」字号应为四号(14pt)，实际为 {font_size}pt")
            if font_name and font_name not in ("SimSun", "宋体", "simsun"):
                issues.append(f"「{kw}」字体应为宋体，实际为 {font_name}")
            if bold:
                issues.append(f"「{kw}」应为非加粗")
            if alignment and alignment not in ("CENTER", "INHERITED"):
                issues.append(f"「{kw}」应居中对齐，实际为 {alignment}")

        if not issues:
            self.report.add_pass("封面信息栏", "信息栏格式符合要求（四号、宋体、非加粗、居中）")
        else:
            self.report.add_fail("封面信息栏", "；".join(issues))

    # ── 4. 封面 —— 导师格式校验 ──

    def _check_cover_advisor_format(self):
        cover_paras = self.parser.get_cover_paragraphs()
        advisor_para = None
        for p in cover_paras:
            if "导师" in p.text or "指导教师" in p.text:
                advisor_para = p
                break

        if advisor_para is None:
            self.report.add_error("封面中未找到导师信息段落")
            return

        text = advisor_para.text.strip()
        colon_pos = max(text.find("："), text.find(":"))
        if colon_pos == -1:
            self.report.add_fail("导师格式", "导师段落未找到冒号分隔符")
            return

        after_colon = text[colon_pos + 1:].strip()

        _tutor_fix_msg = (
            "导师姓名与职称之间必须使用【全角空格】（宽度等同于一个汉字）。\n"
            "检测到您可能使用了普通的半角空格或未加空格。\n"
            "💡 修复建议：请在电脑中文输入法状态下，按下 `Shift + 空格键` 切换为全角模式再输入；\n"
            "或者最简单的方法：直接复制右侧括号内的空白符号【　】去替换文档里的空格。"
        )

        # 检查是否包含半角空格（不应该有）
        if " " in after_colon:
            self.report.add_fail("导师格式", _tutor_fix_msg)
            return

        # 检查是否包含全角空格
        if "　" not in after_colon:
            self.report.add_fail("导师格式", _tutor_fix_msg)
            return

        # 检查全角空格数量是否为1
        full_width_spaces = after_colon.count("　")
        if full_width_spaces != 1:
            self.report.add_fail(
                "导师格式",
                f"导师姓名与职称之间应仅包含一个全角空格，当前有 {full_width_spaces} 个，内容: 「{after_colon}」"
            )
            return

        self.report.add_pass("导师格式", "导师姓名与职称之间包含一个全角空格")

    # ── 5. 封面 —— 日期校验 ──

    def _check_cover_date(self):
        cover_paras = self.parser.get_cover_paragraphs()
        full_text = "\n".join(p.text for p in cover_paras)
        match = self.parser.RE_DATE.search(full_text)

        if match:
            year, month, day = match.groups()
            self.report.add_pass(
                "封面日期",
                f"找到合规日期: {year}-{month.zfill(2)}-{day.zfill(2)}"
            )
        else:
            self.report.add_fail(
                "封面日期",
                "未找到符合 YYYY-MM-DD 或 YYYY年MM月DD日 格式的日期"
            )

    # ── 6. 二级标题结构校验 ──

    def _check_h2_structure(self):
        required_h2 = [
            "1.1 选题目的和意义",
            "1.2 国内外研究现状",
            "1.3 发展趋势",
            "1.4 对本人研究课题的启发",
            "2.1 研究的主要内容",
            "2.2 需要解决的关键问题",
            "2.3 研究的思路",
            "2.4 研究的方法及手段",
        ]
        found_h2 = self.parser.find_h2_titles()
        found_set = set()
        for title in found_h2:
            clean = title.replace(" ", "").replace("　", "")
            for req in required_h2:
                req_clean = req.replace(" ", "").replace("　", "")
                if req_clean in clean or clean in req_clean:
                    found_set.add(req)

        missing = [h for h in required_h2 if h not in found_set]

        if not missing:
            self.report.add_pass("大纲结构", "全部 8 个二级标题均存在")
        else:
            context_lines = []
            if found_h2:
                context_lines.append("文档中实际找到的二级标题：")
                for t in found_h2:
                    context_lines.append(f"  · {t}")
            else:
                context_lines.append("文档中未找到任何二级标题。")
            self.report.add_fail(
                "大纲结构",
                f"缺少以下二级标题: {missing}",
                context_text="\n".join(context_lines)
            )

    # ── 6b. 一级标题结构校验 ──

    def _check_h1_structure(self):
        """校验四个一级标题是否存在，重点验证第二节的完整长标题"""
        # 锚点：必须包含的关键字（去空格后匹配）
        required_h1 = {
            "1. 文献综述": ["文献综述"],
            "2. 研究的主要内容、需要解决的关键问题和思路、研究方法及手段": [
                "主要内容", "关键问题", "思路", "研究方法"
            ],
            "3. 进度安排": ["进度", "安排"],
            "4. 指导教师意见": ["指导教师", "意见"],
        }

        # 提取实际的一级标题
        actual_h1 = []
        for p in self.parser.paragraphs:
            text = p.text.strip()
            if self.parser.RE_H1.match(text) and not self.parser.RE_H2.match(text):
                actual_h1.append(text)

        missing = []
        for expected, keywords in required_h1.items():
            found = False
            for title in actual_h1:
                clean = title.replace(" ", "").replace("　", "")
                if all(kw in clean for kw in keywords):
                    found = True
                    break
            if not found:
                missing.append(expected)

        if not missing:
            self.report.add_pass("一级标题结构", "全部 4 个一级标题均存在且内容完整")
        else:
            context_lines = ["文档中实际找到的一级标题："]
            for t in actual_h1:
                context_lines.append(f"  · {t}")
            self.report.add_fail(
                "一级标题结构",
                f"缺少或内容不完整的标题: {missing}",
                context_text="\n".join(context_lines)
            )

    # ── 7. 文献综述字数校验 ──

    def _check_literature_review_word_count(self):
        word_count = self.parser.count_section_words("1. 文献综述", "2. 研究的主要内容")
        if word_count == 0:
            self.report.add_error("无法定位「1. 文献综述」区域，字数统计跳过")
            return

        min_w = self.config.get("min_words", 1800)
        max_w = self.config.get("max_words", 2200)

        if min_w <= word_count <= max_w:
            self.report.add_pass(
                "文献综述字数",
                f"字数为 {word_count}，在 {min_w}-{max_w} 范围内"
            )
        else:
            self.report.add_fail(
                "文献综述字数",
                f"字数应为 {min_w}-{max_w}，当前为 {word_count} 字"
            )

    # ── 8. 参考文献复杂比例校验（核心算法）──

    def _check_references(self):
        refs = self.parser.get_reference_paragraphs()
        if not refs:
            self.report.add_error("未找到参考文献区域或参考文献为空")
            return

        total = len(refs)

        # 8.1 总数校验
        min_refs = self.config.get("min_refs", 10)
        if total >= min_refs:
            self.report.add_pass("参考文献总数", f"共 {total} 篇，满足 ≥{min_refs} 篇要求")
        else:
            self.report.add_fail("参考文献总数", f"参考文献总数不足 {min_refs} 篇，当前为 {total} 篇")

        # 8.2 年份校验：近3年占比 >= 阈值（严格锁定为当前年-2到当前年）
        current_year = datetime.now().year
        recent_year_start = current_year - 2  # 2026年时为2024
        min_recent_ratio = self.config.get("min_recent_ratio", 1 / 3)
        year_pattern = re.compile(r"(19\d{2}|20\d{2})")
        recent_count = 0
        year_found = 0
        old_refs = []  # 年份不合规的文献
        for idx, ref in enumerate(refs):
            years = year_pattern.findall(ref)
            if years:
                year_found += 1
                latest = max(int(y) for y in years)
                if recent_year_start <= latest <= current_year:
                    recent_count += 1
                else:
                    old_refs.append(f"[{idx+1}] {ref}")

        if year_found > 0:
            ratio = recent_count / year_found
            if ratio >= min_recent_ratio:
                self.report.add_pass(
                    "参考文献年份",
                    f"近 3 年(≥{recent_year_start})文献 {recent_count}/{year_found} 篇，"
                    f"占比 {ratio:.1%}，满足 ≥{min_recent_ratio:.0%} 要求"
                )
            else:
                self.report.add_fail(
                    "参考文献年份",
                    f"近 3 年(≥{recent_year_start})文献 {recent_count}/{year_found} 篇，"
                    f"占比 {ratio:.1%}，不满足 ≥{min_recent_ratio:.0%} 要求",
                    context_text="\n".join(old_refs) if old_refs else None
                )
        else:
            self.report.add_error("参考文献中未能提取到有效年份信息")

        # 8.3 外文文献校验：英文字符占比 > 50% 记为外文，占比 >= 阈值
        min_foreign_ratio = self.config.get("min_foreign_ratio", 1 / 5)
        foreign_count = 0
        non_foreign_refs = []  # 纯中文文献（用于提示外文不足时）
        for idx, ref in enumerate(refs):
            alpha_chars = [c for c in ref if c.isalpha()]
            if alpha_chars:
                english_chars = [c for c in alpha_chars if ord(c) < 128]
                if len(english_chars) / len(alpha_chars) > self.FOREIGN_LANG_RATIO:
                    foreign_count += 1
                else:
                    non_foreign_refs.append(f"[{idx+1}] {ref}")
            else:
                non_foreign_refs.append(f"[{idx+1}] {ref}")

        foreign_ratio = foreign_count / total
        if foreign_ratio >= min_foreign_ratio:
            self.report.add_pass(
                "外文文献比例",
                f"外文文献 {foreign_count}/{total} 篇，占比 {foreign_ratio:.1%}，"
                f"满足 ≥{min_foreign_ratio:.0%} 要求"
            )
        else:
            self.report.add_fail(
                "外文文献比例",
                f"外文文献 {foreign_count}/{total} 篇，占比 {foreign_ratio:.1%}，"
                f"不满足 ≥{min_foreign_ratio:.0%} 要求",
                context_text=f"以下为非外文文献（共 {len(non_foreign_refs)} 篇）：\n" +
                             "\n".join(non_foreign_refs[:10]) if non_foreign_refs else None
            )

        # 8.4 学位论文校验：含 [D] 标识，占比 <= 阈值
        max_thesis_ratio = self.config.get("max_thesis_ratio", 1 / 5)
        thesis_refs = [f"[{idx+1}] {ref}" for idx, ref in enumerate(refs) if self.THESIS_TAG in ref]
        thesis_count = len(thesis_refs)
        thesis_ratio = thesis_count / total
        if thesis_ratio <= max_thesis_ratio:
            self.report.add_pass(
                "学位论文比例",
                f"学位论文 {thesis_count}/{total} 篇，占比 {thesis_ratio:.1%}，"
                f"满足 ≤{max_thesis_ratio:.0%} 要求"
            )
        else:
            self.report.add_fail(
                "学位论文比例",
                f"学位论文 {thesis_count}/{total} 篇，占比 {thesis_ratio:.1%}，"
                f"不满足 ≤{max_thesis_ratio:.0%} 要求",
                context_text="\n".join(thesis_refs) if thesis_refs else None
            )

    # ── 9. 指导教师意见关键字校验 ──

    def _check_supervisor_opinion(self):
        idx = self.parser.find_paragraph_index("4. 指导教师意见")
        if idx is None:
            self.report.add_error("未找到「4. 指导教师意见」部分")
            return

        opinion_text = ""
        for i in range(idx + 1, len(self.parser.paragraphs)):
            p = self.parser.paragraphs[i]
            if self.parser.RE_H1.match(p.text.strip()):
                break
            opinion_text += p.text

        if "同意开题" in opinion_text:
            self.report.add_pass("指导教师意见", "包含「同意开题」关键字")
        elif "不同意开题" in opinion_text:
            self.report.add_fail("指导教师意见", "包含「不同意开题」关键字")
        else:
            self.report.add_fail(
                "指导教师意见",
                "未找到「同意开题」或「不同意开题」关键字"
            )

    # ── 10. 引用一致性双向交叉比对 ──

    def _check_citation_consistency(self):
        """
        双向比对正文引用 [n] 与文末参考文献列表 [n]，
        检测「幽灵引用」和「孤儿文献」。
        """
        body_set = self.parser.extract_body_citation_indices()
        ref_set = self.parser.extract_ref_list_indices()

        # 如果参考文献区域为空或正文无引用，跳过
        if not ref_set and not body_set:
            self.report.add_error("正文与参考文献均未检测到引用标记，跳过一致性比对")
            return
        if not ref_set:
            self.report.add_error("参考文献列表为空，无法进行一致性比对")
            return
        if not body_set:
            self.report.add_error("正文中未检测到任何 [n] 引用标记，无法进行一致性比对")
            return

        # 检查项 A：正文引用了但文献列表没有（幽灵引用）
        ghost = sorted(body_set - ref_set)
        if ghost:
            self.report.add_fail(
                "引用一致性(正向)",
                f"正文引用了文献序号 {ghost}，但文末参考文献列表中并未列出",
                context_text=f"正文引用但文献列表缺失的序号: {ghost}"
            )

        # 检查项 B：文献列表写了但正文从未引用（孤儿文献）
        unreferenced = sorted(ref_set - body_set)
        if unreferenced:
            self.report.add_fail(
                "引用一致性(反向)",
                f"文末列出了文献序号 {unreferenced}，但在正文中从未被引用标记",
                context_text=f"文献列表存在但正文未引用的序号: {unreferenced}"
            )

        # 完全一致
        if not ghost and not unreferenced:
            self.report.add_pass(
                "引用一致性",
                f"正文引用标注与文末参考文献列表完全双向一致（共 {len(ref_set)} 篇）"
            )

    # ── 11. 参考文献上方空行检查 ──

    def _check_ref_preceding_blank(self):
        """检查参考文献标题前是否有空行"""
        ref_idx = None
        for i, p in enumerate(self.parser.paragraphs):
            text = p.text.strip()
            if text == "参考文献" and not self.parser.RE_H2.match(text):
                ref_idx = i
                break

        if ref_idx is None:
            self.report.add_error("未找到「参考文献」标题，跳过空行检查")
            return

        if ref_idx == 0:
            self.report.add_fail(
                "参考文献空行",
                "「参考文献」标题位于文档开头，前面没有空行"
            )
            return

        prev_para = self.parser.paragraphs[ref_idx - 1]
        prev_text = prev_para.text.strip()

        if prev_text:
            self.report.add_fail(
                "参考文献空行",
                f"「参考文献」标题前应有空行，但前一段落不为空: 「{prev_text[:50]}」",
                context_text=f"前一段落内容: {prev_text}"
            )
        else:
            self.report.add_pass("参考文献空行", "「参考文献」标题前已空一行")

    # ── 12. 参考文献悬挂缩进检查 ──

    def _check_ref_hanging_indent(self):
        """检查参考文献条目是否设置了悬挂缩进（first_line_indent < 0 表示悬挂缩进）"""
        ref_start_idx = None
        for i, p in enumerate(self.parser.paragraphs):
            if "参考文献" in p.text and not self.parser.RE_H2.match(p.text.strip()):
                ref_start_idx = i
                break

        if ref_start_idx is None:
            self.report.add_error("未找到参考文献区域，跳过悬挂缩进检查")
            return

        no_indent_refs = []
        count = 0
        for i in range(ref_start_idx + 1, len(self.parser.paragraphs)):
            p = self.parser.paragraphs[i]
            text = p.text.strip()
            if self.parser.RE_H1.match(text):
                break
            if not text:
                continue
            count += 1
            if count > 30:  # 只检查前30条
                break

            # 检查悬挂缩进：first_line_indent 为负值表示悬挂缩进
            pf = p.paragraph_format
            first_indent = pf.first_line_indent
            has_hanging = first_indent is not None and first_indent < 0

            if not has_hanging:
                snippet = text[:50] + ("..." if len(text) > 50 else "")
                no_indent_refs.append(f"[{count}] {snippet}")

        if not no_indent_refs:
            self.report.add_pass("参考文献缩进", "全部参考文献条目已设置悬挂缩进")
        else:
            self.report.add_fail(
                "参考文献缩进",
                f"共 {len(no_indent_refs)} 条参考文献未设置悬挂缩进",
                context_text="\n".join(no_indent_refs[:10])
            )

    # ── 13. 进度安排时间线按序检查 ──

    def _check_schedule_timeline(self):
        """检查第3节进度安排中的日期是否按升序排列"""
        # 找到第3节
        section3_start = None
        section3_end = None
        for i, p in enumerate(self.parser.paragraphs):
            text = p.text.strip()
            if re.match(r"^\s*3\.\s*", text) and ("进度" in text and "安排" in text):
                section3_start = i
                continue
            if section3_start is not None:
                if self.parser.RE_H1.match(text) and i > section3_start:
                    section3_end = i
                    break

        if section3_start is None:
            self.report.add_error("未找到第3节（进度安排），跳过时间线检查")
            return

        if section3_end is None:
            section3_end = len(self.parser.paragraphs)

        # 提取日期
        date_pattern = re.compile(r"(\d{4})\s*[-年]\s*(\d{1,2})\s*[-月]\s*(\d{1,2})\s*日?")
        dates_found = []
        for i in range(section3_start, section3_end):
            p = self.parser.paragraphs[i]
            text = p.text.strip()
            matches = date_pattern.findall(text)
            for match in matches:
                year, month, day = match
                try:
                    date_obj = datetime(int(year), int(month), int(day))
                    dates_found.append((date_obj, text[:50]))
                except ValueError:
                    pass

        if len(dates_found) < 2:
            self.report.add_error("进度安排中日期不足2个，无法检查时间顺序")
            return

        # 检查是否严格升序
        out_of_order = []
        for i in range(1, len(dates_found)):
            if dates_found[i][0] < dates_found[i-1][0]:
                out_of_order.append(
                    f"第{i}项({dates_found[i-1][0].strftime('%Y-%m-%d')}) "
                    f"晚于第{i+1}项({dates_found[i][0].strftime('%Y-%m-%d')})"
                )

        if not out_of_order:
            self.report.add_pass("进度安排", f"共 {len(dates_found)} 个日期，全部按时间升序排列")
        else:
            context_lines = []
            for i, (date, snippet) in enumerate(dates_found):
                context_lines.append(f"[{i+1}] {date.strftime('%Y-%m-%d')} - {snippet}")
            self.report.add_fail(
                "进度安排",
                f"进度安排中存在 {len(out_of_order)} 处日期顺序错误",
                context_text="\n".join(context_lines)
            )

    # ── 14. 双导师签名顺序检查 ──

    def _check_dual_advisor_order(self):
        """检查第4节中双导师签名顺序（校内导师应在校外导师前面）"""
        idx = self.parser.find_paragraph_index("4. 指导教师意见")
        if idx is None:
            self.report.add_error("未找到「4. 指导教师意见」部分，跳过双导师检查")
            return

        # 收集第4节的所有文本
        section_text = ""
        for i in range(idx + 1, len(self.parser.paragraphs)):
            p = self.parser.paragraphs[i]
            if self.parser.RE_H1.match(p.text.strip()):
                break
            section_text += p.text + "\n"

        # 检查是否涉及双导师
        has_dual = "双导师" in section_text

        # 查找校内/校外导师的提及
        internal_pos = -1
        external_pos = -1

        # 查找"校内导师"或"第一导师"的位置
        for kw in ["校内导师", "第一导师", "校内指导教师"]:
            pos = section_text.find(kw)
            if pos != -1:
                internal_pos = pos
                break

        # 查找"校外导师"或"企业导师"或"第二导师"的位置
        for kw in ["校外导师", "企业导师", "第二导师", "校外指导教师"]:
            pos = section_text.find(kw)
            if pos != -1:
                external_pos = pos
                break

        # 如果没有明确的双导师标识，跳过
        if not has_dual and internal_pos == -1 and external_pos == -1:
            return

        # 如果只提到了一个导师，跳过
        if internal_pos == -1 or external_pos == -1:
            self.report.add_pass("双导师顺序", "未检测到双导师配置或仅提及一位导师")
            return

        # 检查校内导师是否在校外导师前面
        if internal_pos < external_pos:
            self.report.add_pass("双导师顺序", "校内导师排在校外导师前面，符合要求")
        else:
            self.report.add_fail(
                "双导师顺序",
                "校内导师应排在校外导师前面，当前顺序不符",
                context_text=f"校内导师位置: 字符{internal_pos}, 校外导师位置: 字符{external_pos}"
            )

    # ── 15. 字体与字号校验 ──

    def _check_fonts_and_sizes(self):
        """
        校验文档 6 个区域的字体与字号：
        1. 封面主标题（宋体、一号、加粗）
        2. 封面信息栏（宋体、四号、非加粗）
        3. 二级标题（宋体、小四、加粗）
        4. 正文（宋体、小四、非加粗）
        5. 参考文献标题（黑体、五号）
        6. 参考文献条目（中文=宋体/五号，英文=Times New Roman/五号）
        """
        FONT_SIZE_MAP = self.parser.FONT_SIZE_MAP
        PT_TOLERANCE = 1.0  # 字号容差 ±1pt

        def _font_ok(actual: Optional[str], expected_list: List[str]) -> bool:
            """判断字体是否在预期列表中（不区分大小写）"""
            if actual is None:
                return True  # 无法检测时跳过
            return any(e.lower() in actual.lower() for e in expected_list)

        def _size_ok(actual: Optional[float], expected_pt: float) -> bool:
            """判断字号是否在容差范围内"""
            if actual is None:
                return True
            return abs(actual - expected_pt) <= PT_TOLERANCE

        def _format_font_info(para) -> str:
            """格式化段落的字体信息，用于 context_text"""
            font = self.parser.get_paragraph_effective_font(para)
            size = self.parser.get_paragraph_font_size(para)
            bold = self.parser.get_paragraph_bold(para)
            parts = []
            if font:
                parts.append(f"字体: {font}")
            if size:
                parts.append(f"字号: {size}pt")
            if bold is not None:
                parts.append(f"加粗: {'是' if bold else '否'}")
            return ", ".join(parts) if parts else "无法检测格式信息"

        # ── 11.1 封面主标题 ──
        cover_paras = self.parser.get_cover_paragraphs()
        target_text = "毕业论文(设计)开题报告"
        target_para = None
        for p in cover_paras:
            clean = p.text.replace("（", "(").replace("）", ")").strip()
            if target_text in clean or ("毕业论文" in clean and "开题报告" in clean):
                target_para = p
                break

        if target_para is None:
            self.report.add_error("封面主标题未找到，跳过字体校验")
        else:
            font = self.parser.get_paragraph_effective_font(target_para)
            size = self.parser.get_paragraph_font_size(target_para)
            bold = self.parser.get_paragraph_bold(target_para)
            issues = []
            if not _font_ok(font, ["SimSun", "宋体"]):
                issues.append(f"字体应为宋体，实际为 {font}")
            if not _size_ok(size, FONT_SIZE_MAP["一号"]):
                issues.append(f"字号应为一号(36pt)，实际为 {size}pt")
            if bold is not None and not bold:
                issues.append("应为加粗，实际为非加粗")

            if not issues:
                self.report.add_pass("封面主标题字体", "主标题格式符合要求（宋体、一号、加粗）")
            else:
                self.report.add_fail(
                    "封面主标题字体",
                    "；".join(issues),
                    context_text=_format_font_info(target_para)
                )

        # ── 11.2 封面信息栏 ──
        info_keywords = ["题目", "学院", "专业", "班级", "学号", "姓名", "导师"]
        info_paras = []
        for p in cover_paras:
            text = p.text.strip()
            for kw in info_keywords:
                if kw in text and len(text) < 80:
                    info_paras.append((kw, p))
                    break

        if not info_paras:
            self.report.add_error("封面信息栏未找到，跳过字体校验")
        else:
            info_issues = []
            info_contexts = []
            for kw, p in info_paras:
                font = self.parser.get_paragraph_effective_font(p)
                size = self.parser.get_paragraph_font_size(p)
                bold = self.parser.get_paragraph_bold(p)
                p_issues = []
                if not _font_ok(font, ["SimSun", "宋体"]):
                    p_issues.append(f"字体应为宋体，实际为 {font}")
                if not _size_ok(size, FONT_SIZE_MAP["四号"]):
                    p_issues.append(f"字号应为四号(14pt)，实际为 {size}pt")
                if bold:
                    p_issues.append("应为非加粗，实际为加粗")
                if p_issues:
                    info_issues.append(f"「{kw}」" + "；".join(p_issues))
                    info_contexts.append(f"[{kw}] {_format_font_info(p)}")

            if not info_issues:
                self.report.add_pass("封面信息栏字体", "信息栏格式符合要求（宋体、四号、非加粗）")
            else:
                self.report.add_fail(
                    "封面信息栏字体",
                    "；".join(info_issues),
                    context_text="\n".join(info_contexts)
                )

        # ── 11.3 二级标题 ──
        h2_paras = [p for p in self.parser.paragraphs if self.parser.RE_H2.match(p.text.strip())]
        if not h2_paras:
            self.report.add_error("未找到二级标题，跳过字体校验")
        else:
            h2_issues = []
            h2_contexts = []
            for p in h2_paras:
                title = p.text.strip()
                font = self.parser.get_paragraph_effective_font(p)
                size = self.parser.get_paragraph_font_size(p)
                bold = self.parser.get_paragraph_bold(p)
                p_issues = []
                if not _font_ok(font, ["SimSun", "宋体"]):
                    p_issues.append(f"字体应为宋体，实际为 {font}")
                if not _size_ok(size, FONT_SIZE_MAP["小四"]):
                    p_issues.append(f"字号应为小四(12pt)，实际为 {size}pt")
                if bold is not None and not bold:
                    p_issues.append("应为加粗，实际为非加粗")
                if p_issues:
                    h2_issues.append(f"「{title}」" + "；".join(p_issues))
                    h2_contexts.append(f"[{title}] {_format_font_info(p)}")

            if not h2_issues:
                self.report.add_pass("二级标题字体", "全部二级标题格式符合要求（宋体、小四、加粗）")
            else:
                self.report.add_fail(
                    "二级标题字体",
                    f"共 {len(h2_issues)} 个二级标题格式不合规",
                    context_text="\n".join(h2_contexts[:10])
                )

        # ── 11.4 正文 ──
        cover_paras = self.parser.get_cover_paragraphs()
        body_paras = self.parser.paragraphs[len(cover_paras):]
        # 排除二级标题、参考文献区域
        body_text_paras = []
        in_refs = False
        for p in body_paras:
            text = p.text.strip()
            if "参考文献" in text and not self.parser.RE_H2.match(text):
                in_refs = True
                continue
            if in_refs:
                continue
            if self.parser.RE_H2.match(text):
                continue
            if text and len(text) > 20:  # 只检查有意义的正文段落
                body_text_paras.append(p)

        if not body_text_paras:
            self.report.add_error("未找到正文段落，跳过字体校验")
        else:
            body_issues = []
            body_contexts = []
            for p in body_text_paras[:20]:  # 只检查前20段，避免过多输出
                font = self.parser.get_paragraph_effective_font(p)
                size = self.parser.get_paragraph_font_size(p)
                bold = self.parser.get_paragraph_bold(p)
                p_issues = []
                if not _font_ok(font, ["SimSun", "宋体"]):
                    p_issues.append(f"字体应为宋体，实际为 {font}")
                if not _size_ok(size, FONT_SIZE_MAP["小四"]):
                    p_issues.append(f"字号应为小四(12pt)，实际为 {size}pt")
                if bold:
                    p_issues.append("应为非加粗，实际为加粗")
                if p_issues:
                    snippet = p.text[:30] + ("..." if len(p.text) > 30 else "")
                    body_issues.append(f"「{snippet}」" + "；".join(p_issues))
                    body_contexts.append(f"[{snippet}] {_format_font_info(p)}")

            if not body_issues:
                self.report.add_pass("正文字体", "正文格式符合要求（宋体、小四、非加粗）")
            else:
                self.report.add_fail(
                    "正文字体",
                    f"共 {len(body_issues)} 个正文段落格式不合规",
                    context_text="\n".join(body_contexts[:10])
                )

        # ── 11.5 参考文献标题 ──
        ref_title_para = None
        for p in self.parser.paragraphs:
            text = p.text.strip()
            if "参考文献" in text and not self.parser.RE_H2.match(text):
                ref_title_para = p
                break
        if ref_title_para is None:
            self.report.add_error("未找到独立的「参考文献」标题段落，跳过字体校验")
        else:
            font = self.parser.get_paragraph_effective_font(ref_title_para)
            size = self.parser.get_paragraph_font_size(ref_title_para)
            issues = []
            if not _font_ok(font, ["SimHei", "黑体"]):
                issues.append(f"字体应为黑体，实际为 {font}")
            if not _size_ok(size, FONT_SIZE_MAP["五号"]):
                issues.append(f"字号应为五号(10.5pt)，实际为 {size}pt")

            if not issues:
                self.report.add_pass("参考文献标题字体", "「参考文献」标题格式符合要求（黑体、五号）")
            else:
                self.report.add_fail(
                    "参考文献标题字体",
                    "；".join(issues),
                    context_text=_format_font_info(ref_title_para)
                )

        # ── 11.6 参考文献条目 ──
        ref_paras = self.parser.get_reference_paragraphs()
        if not ref_paras:
            self.report.add_error("未找到参考文献条目，跳过字体校验")
        else:
            # 获取参考文献区域的段落对象
            ref_start_idx = None
            for i, p in enumerate(self.parser.paragraphs):
                if "参考文献" in p.text and not self.parser.RE_H2.match(p.text.strip()):
                    ref_start_idx = i
                    break

            if ref_start_idx is not None:
                ref_item_issues = []
                ref_item_contexts = []
                count = 0
                for i in range(ref_start_idx + 1, len(self.parser.paragraphs)):
                    p = self.parser.paragraphs[i]
                    text = p.text.strip()
                    if self.parser.RE_H1.match(text):
                        break
                    if not text:
                        continue
                    count += 1
                    if count > 20:  # 只检查前20条
                        break

                    # 判断是否为外文文献（包含中文字符则为中文文献）
                    chinese_chars = [c for c in text if '一' <= c <= '鿿']
                    is_foreign = len(chinese_chars) == 0

                    font = self.parser.get_paragraph_effective_font(p)
                    size = self.parser.get_paragraph_font_size(p)
                    p_issues = []

                    if is_foreign:
                        if not _font_ok(font, ["Times New Roman"]):
                            p_issues.append(f"英文字体应为Times New Roman，实际为 {font}")
                    else:
                        if not _font_ok(font, ["SimSun", "宋体"]):
                            p_issues.append(f"中文字体应为宋体，实际为 {font}")

                    if not _size_ok(size, FONT_SIZE_MAP["五号"]):
                        p_issues.append(f"字号应为五号(10.5pt)，实际为 {size}pt")

                    if p_issues:
                        snippet = text[:40] + ("..." if len(text) > 40 else "")
                        ref_item_issues.append(f"[{count}] {snippet}：" + "；".join(p_issues))
                        ref_item_contexts.append(f"[{count}] {_format_font_info(p)}")

                if not ref_item_issues:
                    self.report.add_pass("参考文献条目字体", "参考文献条目格式符合要求（中文宋体/英文Times New Roman、五号）")
                else:
                    self.report.add_fail(
                        "参考文献条目字体",
                        f"共 {len(ref_item_issues)} 条参考文献格式不合规",
                        context_text="\n".join(ref_item_contexts[:10])
                    )


# ─────────────────────────────────────────────────────────────
# HTMLReporter —— 生成精美的单文件 HTML 诊断报告
# ─────────────────────────────────────────────────────────────

class HTMLReporter:
    """
    职责：接收 AnalysisReport 的结果字典，渲染为一个自包含的 HTML 文件，
    使用现代卡片式设计，无外部依赖。
    """

    def __init__(self, result: Dict[str, Any], doc_name: str = ""):
        self.result = result
        self.doc_name = doc_name
        self.scan_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        self.summary = result.get("summary", {})

    def generate(self, output_path: str = "report.html") -> str:
        """生成 HTML 文件并返回其绝对路径"""
        html = self._render()
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html)
        return os.path.abspath(output_path)

    def _render(self) -> str:
        total = self.summary.get("total_checks", 0)
        passed = self.summary.get("passed", 0)
        failed = self.summary.get("failed", 0)
        errors = self.summary.get("errors", 0)

        passed_cards = self._render_cards(self.result.get("passed_items", []), is_pass=True)
        failed_cards = self._render_cards(self.result.get("failed_items", []), is_pass=False)
        error_list = self._render_errors(self.result.get("errors", []))

        # 计算通过率百分比
        pass_rate = f"{passed / total * 100:.0f}%" if total > 0 else "N/A"

        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>开题报告格式校验诊断报告</title>
<style>
  :root {{
    --green: #22c55e;
    --green-bg: #f0fdf4;
    --green-border: #bbf7d0;
    --red: #ef4444;
    --red-bg: #fef2f2;
    --red-border: #fecaca;
    --orange: #f97316;
    --orange-bg: #fff7ed;
    --orange-border: #fed7aa;
    --gray-50: #f9fafb;
    --gray-100: #f3f4f6;
    --gray-200: #e5e7eb;
    --gray-300: #d1d5db;
    --gray-600: #4b5563;
    --gray-800: #1f2937;
    --gray-900: #111827;
    --radius: 12px;
    --shadow: 0 1px 3px rgba(0,0,0,0.08), 0 1px 2px rgba(0,0,0,0.06);
    --shadow-md: 0 4px 6px rgba(0,0,0,0.07), 0 2px 4px rgba(0,0,0,0.06);
    --shadow-lg: 0 10px 15px rgba(0,0,0,0.1), 0 4px 6px rgba(0,0,0,0.05);
  }}
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, "Noto Sans SC", sans-serif;
    background: linear-gradient(135deg, #f0f4ff 0%, #faf5ff 50%, #fef2f2 100%);
    color: var(--gray-800);
    min-height: 100vh;
    padding: 0 0 60px 0;
  }}

  /* ── Header ── */
  .header {{
    background: linear-gradient(135deg, #1e293b 0%, #334155 100%);
    color: #fff;
    padding: 40px 24px 32px;
    text-align: center;
    box-shadow: var(--shadow-lg);
  }}
  .header h1 {{
    font-size: 28px;
    font-weight: 700;
    letter-spacing: 1px;
    margin-bottom: 8px;
  }}
  .header .subtitle {{
    font-size: 14px;
    color: #94a3b8;
  }}
  .header .meta {{
    margin-top: 16px;
    display: flex;
    justify-content: center;
    gap: 24px;
    font-size: 13px;
    color: #cbd5e1;
  }}
  .header .meta span {{
    display: flex;
    align-items: center;
    gap: 6px;
  }}

  /* ── Container ── */
  .container {{
    max-width: 960px;
    margin: 0 auto;
    padding: 0 20px;
  }}

  /* ── Dashboard ── */
  .dashboard {{
    display: grid;
    grid-template-columns: repeat(4, 1fr);
    gap: 16px;
    margin: -30px 0 32px;
    position: relative;
    z-index: 10;
  }}
  .stat-card {{
    background: #fff;
    border-radius: var(--radius);
    padding: 24px 16px;
    text-align: center;
    box-shadow: var(--shadow-md);
    border-top: 4px solid transparent;
    transition: transform 0.2s;
  }}
  .stat-card:hover {{ transform: translateY(-2px); box-shadow: var(--shadow-lg); }}
  .stat-card .number {{
    font-size: 36px;
    font-weight: 800;
    line-height: 1;
    margin-bottom: 6px;
  }}
  .stat-card .label {{
    font-size: 13px;
    color: var(--gray-600);
    font-weight: 500;
  }}
  .stat-card.total {{ border-top-color: #6366f1; }}
  .stat-card.total .number {{ color: #6366f1; }}
  .stat-card.passed {{ border-top-color: var(--green); }}
  .stat-card.passed .number {{ color: var(--green); }}
  .stat-card.failed {{ border-top-color: var(--red); }}
  .stat-card.failed .number {{ color: var(--red); }}
  .stat-card.errors {{ border-top-color: var(--orange); }}
  .stat-card.errors .number {{ color: var(--orange); }}

  /* ── Pass Rate Bar ── */
  .rate-bar-wrap {{
    background: #fff;
    border-radius: var(--radius);
    padding: 20px 24px;
    margin-bottom: 32px;
    box-shadow: var(--shadow);
  }}
  .rate-bar-wrap .rate-label {{
    display: flex;
    justify-content: space-between;
    margin-bottom: 10px;
    font-size: 14px;
    font-weight: 600;
    color: var(--gray-600);
  }}
  .rate-bar {{
    width: 100%;
    height: 12px;
    background: var(--gray-200);
    border-radius: 6px;
    overflow: hidden;
  }}
  .rate-bar .fill {{
    height: 100%;
    border-radius: 6px;
    background: linear-gradient(90deg, #22c55e, #4ade80);
    transition: width 0.6s ease;
  }}

  /* ── Section Titles ── */
  .section-title {{
    font-size: 20px;
    font-weight: 700;
    margin: 32px 0 16px;
    padding-left: 14px;
    border-left: 4px solid;
    line-height: 1.3;
  }}
  .section-title.pass {{ border-left-color: var(--green); color: var(--gray-800); }}
  .section-title.fail {{ border-left-color: var(--red); color: var(--gray-800); }}
  .section-title.err  {{ border-left-color: var(--orange); color: var(--gray-800); }}
  .section-title .count {{
    font-size: 14px;
    font-weight: 500;
    color: var(--gray-600);
    margin-left: 8px;
  }}

  /* ── Cards ── */
  .card-grid {{
    display: flex;
    flex-direction: column;
    gap: 12px;
  }}
  .card {{
    background: #fff;
    border-radius: var(--radius);
    padding: 18px 22px;
    box-shadow: var(--shadow);
    border-left: 5px solid transparent;
    transition: box-shadow 0.2s;
    display: flex;
    align-items: flex-start;
    gap: 14px;
  }}
  .card:hover {{ box-shadow: var(--shadow-md); }}
  .card.pass {{ border-left-color: var(--green); background: var(--green-bg); }}
  .card.fail {{ border-left-color: var(--red); background: var(--red-bg); }}
  .card .icon {{
    font-size: 22px;
    flex-shrink: 0;
    margin-top: 1px;
  }}
  .card .content {{ flex: 1; }}
  .card .card-name {{
    font-weight: 700;
    font-size: 15px;
    margin-bottom: 4px;
  }}
  .card.pass .card-name {{ color: #166534; }}
  .card.fail .card-name {{ color: #991b1b; }}
  .card .card-msg {{
    font-size: 13.5px;
    line-height: 1.6;
    color: var(--gray-600);
  }}
  .card.fail .card-msg {{
    color: #b91c1c;
    font-weight: 500;
  }}

  /* ── Error Context Box ── */
  .error-context-box {{
    background-color: #fef2f2;
    border-left: 4px solid #ef4444;
    padding: 12px 16px;
    margin-top: 10px;
    border-radius: 6px;
    font-family: Consolas, Monaco, "Courier New", monospace;
    font-size: 13px;
    color: #991b1b;
    white-space: pre-wrap;
    word-break: break-all;
    text-align: left;
    line-height: 1.7;
  }}
  .error-context-box .ctx-hint {{
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
    font-weight: 700;
    font-size: 12px;
    color: #b91c1c;
    margin-bottom: 6px;
    letter-spacing: 0.3px;
  }}

  /* ── Error List ── */
  .error-list {{
    display: flex;
    flex-direction: column;
    gap: 8px;
  }}
  .error-item {{
    background: var(--orange-bg);
    border: 1px solid var(--orange-border);
    border-radius: 8px;
    padding: 12px 18px;
    font-size: 13.5px;
    color: #9a3412;
    display: flex;
    align-items: flex-start;
    gap: 10px;
  }}
  .error-item .icon {{ flex-shrink: 0; font-size: 16px; }}

  /* ── Footer ── */
  .footer {{
    text-align: center;
    margin-top: 48px;
    padding: 24px;
    font-size: 12px;
    color: var(--gray-300);
  }}

  /* ── Responsive ── */
  @media (max-width: 640px) {{
    .dashboard {{ grid-template-columns: repeat(2, 1fr); }}
    .header h1 {{ font-size: 22px; }}
    .stat-card .number {{ font-size: 28px; }}
  }}
</style>
</head>
<body>

<!-- Header -->
<div class="header">
  <h1>开题报告格式校验诊断报告</h1>
  <div class="subtitle">Graduation Thesis Proposal Format Verification Report</div>
  <div class="meta">
    <span>&#128196; 文件：{self._escape(self.doc_name or "未指定")}</span>
    <span>&#128339; 扫描时间：{self.scan_time}</span>
  </div>
</div>

<div class="container">

  <!-- Dashboard -->
  <div class="dashboard">
    <div class="stat-card total">
      <div class="number">{total}</div>
      <div class="label">总检查项</div>
    </div>
    <div class="stat-card passed">
      <div class="number">{passed}</div>
      <div class="label">通过</div>
    </div>
    <div class="stat-card failed">
      <div class="number">{failed}</div>
      <div class="label">待修改</div>
    </div>
    <div class="stat-card errors">
      <div class="number">{errors}</div>
      <div class="label">异常</div>
    </div>
  </div>

  <!-- Pass Rate Bar -->
  <div class="rate-bar-wrap">
    <div class="rate-label">
      <span>合规率</span>
      <span>{pass_rate}</span>
    </div>
    <div class="rate-bar">
      <div class="fill" style="width: {pass_rate};"></div>
    </div>
  </div>

  <!-- Failed Items -->
  {self._section("待修改项", failed, "fail", failed_cards)}

  <!-- Passed Items -->
  {self._section("已通过项", passed, "pass", passed_cards)}

  <!-- Errors -->
  {self._render_error_section(errors, error_list)}

</div>

<div class="footer">
  Format Verify Tool &mdash; 开题报告格式自动校验工具 v1.0
</div>

</body>
</html>"""

    def _render_cards(self, items: List[Dict[str, str]], is_pass: bool) -> str:
        if not items:
            return '<div style="color:#9ca3af;font-size:14px;padding:12px 0;">暂无</div>'
        cls = "pass" if is_pass else "fail"
        icon = "&#10003;" if is_pass else "&#10007;"
        cards = []
        for item in items:
            context_html = ""
            ctx = item.get("context_text")
            if ctx:
                context_html = (
                    f'<div class="error-context-box">'
                    f'<div class="ctx-hint">'
                    f'\U0001f50d 引发错误的段落/文献原文：</div>'
                    f'{self._escape(ctx)}</div>'
                )
            cards.append(f"""
    <div class="card {cls}">
      <div class="icon">{icon}</div>
      <div class="content">
        <div class="card-name">{self._escape(item.get('name', ''))}</div>
        <div class="card-msg">{self._escape(item.get('message', ''))}</div>
        {context_html}
      </div>
    </div>""")
        return "\n".join(cards)

    def _render_errors(self, errors: List[str]) -> str:
        if not errors:
            return ""
        items = []
        for err in errors:
            items.append(f"""
    <div class="error-item">
      <span class="icon">&#9888;</span>
      <span>{self._escape(err)}</span>
    </div>""")
        return "\n".join(items)

    def _section(self, title: str, count: int, cls: str, cards_html: str) -> str:
        return f"""
  <div class="section-title {cls}">{title}<span class="count">({count})</span></div>
  <div class="card-grid">
    {cards_html}
  </div>"""

    def _render_error_section(self, count: int, error_html: str) -> str:
        if not error_html:
            return ""
        return f"""
  <div class="section-title err">异常 / 警告<span class="count">({count})</span></div>
  <div class="error-list">
    {error_html}
  </div>"""

    @staticmethod
    def _escape(text: str) -> str:
        """HTML 转义"""
        return (
            text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace('"', "&quot;")
                .replace("'", "&#39;")
        )


# ─────────────────────────────────────────────────────────────
# 主入口函数
# ─────────────────────────────────────────────────────────────

def _default_config() -> Dict[str, Any]:
    """返回一份默认配置，所有阈值与硬编码版本一致"""
    return {
        "line_spacing": 1.5,      # 预期行距倍数
        "min_refs": 10,           # 参考文献最少篇数
        "min_words": 1800,        # 文献综述最少字数
        "max_words": 2200,        # 文献综述最多字数
        "recent_year": 2024,      # "近 N 年"阈值
        "min_recent_ratio": 1/3,  # 近年文献最低占比
        "min_foreign_ratio": 1/5, # 外文文献最低占比
        "max_thesis_ratio": 1/5,  # 学位论文最高占比
    }


def analyze_proposal(file_path: str, config: dict = None) -> Dict[str, Any]:
    """
    校验入口：接收 .docx 文件路径和可选配置字典，返回结构化校验报告字典。

    Args:
        file_path: .docx 文件路径
        config: 可选的自定义阈值配置，缺失的键会用默认值补齐

    Returns:
        {
            "passed_items": [{"name": "...", "message": "..."}, ...],
            "failed_items": [{"name": "...", "message": "..."}, ...],
            "errors": ["..."],
            "summary": {"total_checks": N, "passed": N, "failed": N, "errors": N}
        }
    """
    # 合并用户配置与默认配置
    merged = _default_config()
    if config:
        merged.update({k: v for k, v in config.items() if v is not None})

    report = AnalysisReport()
    try:
        parser = DocumentParser(file_path).load()
    except Exception as e:
        report.add_error(f"文件读取失败: {str(e)}")
        return report.to_dict()

    validator = Validator(parser, config=merged)
    return validator.run_all()


# ─────────────────────────────────────────────────────────────
# CLI 入口
# ─────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: python Format_verify_tool.py <path_to_docx>")
        print("\n--- Mock 数据说明 ---")
        print("请准备一个 .docx 文件作为测试输入，文件应包含：")
        print("  1. 封面页（主标题「毕业论文(设计)开题报告」，信息栏，日期）")
        print("  2. 二级标题：1.1~1.4, 2.1~2.4")
        print("  3. 文献综述区域（1800-2200 字）")
        print("  4. 参考文献区域（≥10 篇，含年份、外文、学位论文标识）")
        print("  5. 指导教师意见区域（含「同意开题」字样）")
        sys.exit(1)

    file_path = sys.argv[1]

    # 1. 运行校验
    result = analyze_proposal(file_path)

    # 2. 生成 HTML 报告
    doc_name = os.path.basename(file_path)
    reporter = HTMLReporter(result, doc_name=doc_name)
    report_path = reporter.generate("report.html")

    # 3. 在浏览器中自动打开
    webbrowser.open("file://" + report_path)

    # 4. 控制台提示
    print("校验完成！正在浏览器中打开诊断报告...")
    print(f"报告路径: {report_path}")
